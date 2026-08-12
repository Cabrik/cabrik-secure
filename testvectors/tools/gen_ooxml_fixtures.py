"""Erzeugt echte Office-Dokumente als Vorlagen fuer die Metadatentests.

Warum echte Dateien und nicht von Hand gebaute:

Die Modultests in `ooxml.rs` bauen ein Dokument aus einzelnen XML-Teilen. Das
prueft die Logik, aber nicht die Wirklichkeit. Word und LibreOffice schreiben
Teile, an die niemand denkt -- `docProps/thumbnail.jpeg`, `customXml/` mit einer
festen GUID, neun `rsid`-Werte aus der Vorlage. Genau diese kamen erst heraus,
als das erste echte Dokument durch den Pruefer lief.

Aufruf:
    python testvectors/tools/gen_ooxml_fixtures.py
"""

from __future__ import annotations

import datetime
import pathlib
import zipfile

import docx
from docx.shared import Inches
from openpyxl import Workbook

ZIEL = pathlib.Path(__file__).resolve().parents[1] / "metadata"


def _stempel(kern) -> None:
    """Setzt die Kerneigenschaften, die Word beim Speichern fuellt."""
    kern.author = "Dr. Anna Beispiel"
    kern.last_modified_by = "Prof. Carl Chef"
    kern.title = "Angebot Projekt Nordstern"
    kern.subject = "Preisverhandlung"
    kern.keywords = "vertraulich, intern"
    kern.comments = "Nicht an den Kunden geben"
    kern.category = "Angebot"
    kern.revision = 17
    kern.created = datetime.datetime(2026, 3, 1, 9, 12, 0)
    kern.modified = datetime.datetime(2026, 3, 2, 17, 45, 0)


def word_schlicht() -> pathlib.Path:
    """Ein Textdokument mit vollstaendig gefuellten Eigenschaften."""
    d = docx.Document()
    d.add_heading("Vertrauliches Angebot", 0)
    d.add_paragraph("Sehr geehrte Damen und Herren, ")
    d.add_paragraph("unser Angebot liegt bei 240.000 Euro.")
    _stempel(d.core_properties)

    ziel = ZIEL / "dokument_mit_metadaten.docx"
    d.save(ziel)
    return ziel


def word_mit_bild() -> pathlib.Path:
    """Ein Dokument mit eingebettetem Bild.

    Das Bild traegt eigenes EXIF -- ein Fall, den v1 gar nicht kannte: Die
    Bereinigung eines Dokuments muss auch die Bilder darin erfassen.
    """
    quelle = ZIEL / "foto_mit_exif.jpg"
    if not quelle.exists():
        raise SystemExit(
            f"{quelle} fehlt. Zuerst gen_metadata_fixtures.py ausfuehren."
        )

    d = docx.Document()
    d.add_paragraph("Bericht mit Lichtbild:")
    d.add_picture(str(quelle), width=Inches(2.0))
    _stempel(d.core_properties)

    ziel = ZIEL / "dokument_mit_bild.docx"
    d.save(ziel)
    return ziel


def tabelle() -> pathlib.Path:
    """Eine Tabelle -- derselbe Behaelter, anderer Inhalt."""
    wb = Workbook()
    blatt = wb.active
    blatt.title = "Kalkulation"
    blatt["A1"] = "Position"
    blatt["B1"] = "Betrag"
    blatt["A2"] = "Entwicklung"
    blatt["B2"] = 240000

    p = wb.properties
    p.creator = "Dr. Anna Beispiel"
    p.lastModifiedBy = "Prof. Carl Chef"
    p.title = "Kalkulation Nordstern"
    p.description = "Interne Fassung, nicht weitergeben"
    p.category = "Kalkulation"
    p.created = datetime.datetime(2026, 3, 1, 9, 12, 0)
    p.modified = datetime.datetime(2026, 3, 2, 17, 45, 0)

    ziel = ZIEL / "tabelle_mit_metadaten.xlsx"
    wb.save(ziel)
    return ziel


def odf_text() -> pathlib.Path:
    """Ein ODF-Textdokument.

    ODF fuehrt zwei Angaben, die es in OOXML so nicht gibt: die
    Gesamtbearbeitungszeit und die Zahl der Speichervorgaenge.
    """
    from odf import dc
    from odf import meta as odfmeta
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    d = OpenDocumentText()
    d.text.addElement(P(text="Sehr geehrte Damen und Herren, "))
    d.text.addElement(P(text="unser Angebot liegt bei 240.000 Euro."))

    def setze(element, wert):
        e = element()
        e.addText(wert)
        d.meta.addElement(e)

    setze(dc.Creator, "Prof. Carl Chef")
    setze(dc.Title, "Angebot Projekt Nordstern")
    setze(dc.Description, "Nicht an den Kunden geben")
    setze(odfmeta.InitialCreator, "Dr. Anna Beispiel")
    setze(odfmeta.PrintedBy, "Dr. Anna Beispiel")
    setze(odfmeta.EditingDuration, "PT4H12M30S")
    setze(odfmeta.EditingCycles, "23")

    feld = odfmeta.UserDefined(name="Aktenzeichen")
    feld.addText("2026-0815")
    d.meta.addElement(feld)

    ziel = ZIEL / "dokument_mit_metadaten.odt"
    d.save(str(ziel))
    return ziel


def bericht(pfad: pathlib.Path) -> None:
    with zipfile.ZipFile(pfad) as z:
        teile = sorted(z.namelist())
    print(f"\n{pfad.name}  ({pfad.stat().st_size} Bytes, {len(teile)} Teile)")
    for n in teile:
        merk = ""
        if n.startswith("docProps/thumbnail"):
            merk = "   <- Vorschaubild: zweite Kopie des Inhalts"
        elif n.startswith("customXml/"):
            merk = "   <- angehaengtes XML, traegt oft eine feste GUID"
        elif n.endswith("settings.xml"):
            merk = "   <- rsid: Bearbeitungssitzungen"
        print(f"    {n}{merk}")


if __name__ == "__main__":
    ZIEL.mkdir(parents=True, exist_ok=True)
    erzeugt = [word_schlicht(), tabelle(), odf_text()]
    try:
        erzeugt.append(word_mit_bild())
    except SystemExit as e:
        print(f"uebersprungen: {e}")

    for p in erzeugt:
        bericht(p)

    print(
        "\nFertig. Die Rust-Tests unter crates/cabrik-metadata/tests/ooxml_echt.rs\n"
        "lesen diese Dateien und legen ihre Ergebnisse als *.stripped.* daneben."
    )
